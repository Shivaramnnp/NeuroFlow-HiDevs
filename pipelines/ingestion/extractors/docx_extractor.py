from __future__ import annotations

import io
import logging
from typing import Any, Dict, List, Optional, Union

import docx

from .base import BaseExtractor, ExtractedPage
from .pdf_extractor import format_markdown_table

logger = logging.getLogger("neuroflow-docx-extractor")


class DocxExtractor(BaseExtractor):
    """
    DOCX Document Extractor:
    - Uses python-docx
    - Extracts text from paragraphs, table cells, and headers separately
    - Preserves heading hierarchy: {'level': 'h1', 'section': 'Introduction'}
    """

    async def extract(self, source: Union[str, bytes, io.BytesIO], **kwargs) -> List[ExtractedPage]:
        if isinstance(source, bytes):
            file_obj = io.BytesIO(source)
        elif isinstance(source, io.BytesIO):
            file_obj = source
        else:
            file_obj = source

        doc = docx.Document(file_obj)
        results: List[ExtractedPage] = []
        page_number = 1

        # 1. Extract headers from sections
        for s_idx, section in enumerate(doc.sections):
            header = section.header
            if header:
                header_text = "\n".join(
                    p.text.strip() for p in header.paragraphs if p.text.strip()
                )
                if header_text:
                    results.append(
                        ExtractedPage(
                            page_number=page_number,
                            content=header_text,
                            content_type="text",
                            metadata={"type": "header", "section_idx": s_idx},
                        )
                    )

        # 2. Extract paragraphs and preserve heading hierarchy
        current_hierarchy: Dict[str, str] = {}

        # Iterate through paragraphs
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            style_name = p.style.name.lower() if p.style and p.style.name else ""

            if "heading 1" in style_name or style_name == "title":
                current_hierarchy = {"level": "h1", "section": text}
                results.append(
                    ExtractedPage(
                        page_number=page_number,
                        content=text,
                        content_type="text",
                        metadata=dict(current_hierarchy),
                    )
                )
            elif "heading 2" in style_name:
                current_hierarchy = {"level": "h2", "section": text}
                results.append(
                    ExtractedPage(
                        page_number=page_number,
                        content=text,
                        content_type="text",
                        metadata=dict(current_hierarchy),
                    )
                )
            elif "heading 3" in style_name:
                current_hierarchy = {"level": "h3", "section": text}
                results.append(
                    ExtractedPage(
                        page_number=page_number,
                        content=text,
                        content_type="text",
                        metadata=dict(current_hierarchy),
                    )
                )
            else:
                meta = dict(current_hierarchy) if current_hierarchy else {}
                results.append(
                    ExtractedPage(
                        page_number=page_number,
                        content=text,
                        content_type="text",
                        metadata=meta,
                    )
                )

        # 3. Extract tables
        for t_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_data.append(row_cells)
            
            table_md = format_markdown_table(table_data)
            if table_md.strip():
                results.append(
                    ExtractedPage(
                        page_number=page_number,
                        content=table_md,
                        content_type="table",
                        metadata={"table_index": t_idx, **current_hierarchy},
                    )
                )

        return results
